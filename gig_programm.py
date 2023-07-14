import json
import os

from bs4 import BeautifulSoup
from pathlib import Path
import docx
import openpyxl


class Parse:
    def __init__(self):
        self.results = []
        self.html_text = None
        self.prorgamm_matrix = {}
        self.prorgamms = []
        self.themes = []

    def take_from_text_html(self, file_path):
        with open(file_path, 'r', encoding='UTF-8') as file:
            self.html_text = file.read()
        self.find_info()

    def take_from_dir(self, dir_path):
        [self.results.append(Path(elem).stem) for elem in Path(dir_path).iterdir()]

    def format_text(self, text: str):
        return text.replace('\n', ' ').strip().replace('  ', ' ')

    def take_from_docx(self, docx_path):
        doc = docx.Document(docx_path)
        feel = 7
        programm_name = None
        while True:
            if not doc.paragraphs[feel].text:
                feel += 1
                continue
            programm_name = self.format_text(doc.paragraphs[feel].text)
            break

        if programm_name:
            theme_list = [self.format_text(str(row.text)) for row in doc.tables[0].column_cells(1)]
            # for separates in [str(row.text).split('\n') for row in doc.tables[0].column_cells(1)]:
            #     for l in separates:
            #         if l:
            #             theme_list.append(self.format_text(l))
            theme_list.pop(0)
            theme_list.pop(-1)
            self.prorgamms.append(programm_name)
            self.themes.extend(theme_list)
            self.prorgamm_matrix[programm_name] = theme_list
            return
        raise Exception('Programm name cant be None, check the docx_path')

    def take_from_docxs(self, docxs_dir_path):
        for path in Path(docxs_dir_path).iterdir():
            try:
                self.take_from_docx(path)
            except Exception as ex:
                raise Exception(ex, path)


    def find_info(self):

        if self.html_text:

            soup = BeautifulSoup(self.html_text, 'lxml').find_all('li', {'role': 'treeitem'})
            for el in soup:
                self.results.append(el.text)
        else:
            raise Exception('No text to parse, find the ways to get text')

    def save_program_matrix_to_json(self):
        with open('matrix.json', 'w', encoding="UTF-8") as file:
            json.dump(self.prorgamm_matrix, file)

    def save_program_matrix_to_exel(self):
        wb = openpyxl.Workbook()
        sh = wb.worksheets[0]
        line_0 = ['Программа', *set(self.themes)]

        sh.append(line_0)

        for programm, themes in self.prorgamm_matrix.items():
            line = [programm, *' '*len(set(self.themes))]
            for theme in themes:
                if theme in line_0:
                    line[line_0.index(theme)] = '1'

            sh.append(line)
        wb.save('results.xlsx')


if __name__ == '__main__':
    p = Parse()
    # p.take_from_text_html("C:\\Users\zaitsev_ad\Desktop\воркс.txt")
    #
    # p.take_from_dir("C:\\Users\zaitsev_ad\Documents\ЭЛМК\гигиеническая подготовка и аттестация\Attachments_egerevafguz@mail.ru_2023-05-25_11-37-19")
    p.take_from_docxs("C:\\Users\zaitsev_ad\Documents\ЭЛМК\гигиеническая подготовка и аттестация\Attachments_egerevafguz@mail.ru_2023-05-25_11-37-19")

    print(p.prorgamm_matrix)
    print(len(set(p.prorgamms)))
    print(len(set(p.themes)))

    p.save_program_matrix_to_json()
    p.save_program_matrix_to_exel()
    # for result in p.results:
    #     print(result)
